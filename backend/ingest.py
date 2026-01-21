import os
import shutil
from typing import List, Optional
import docx
from langchain_community.document_loaders import PyPDFLoader, TextLoader
try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except ImportError:
    from langchain.text_splitter import RecursiveCharacterTextSplitter

# Updated imports for modern LangChain
try:
    from langchain_core.documents import Document
except ImportError:
    from langchain.docstore.document import Document

import core
import document_processor  # Import the new smart document processor
import smart_processor  # Import the new dual-brain processor

def load_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return [Document(page_content="\n".join(full_text), metadata={"source": file_path})]

def ingest_document(file_path: str, original_filename: str, summary: str = "") -> List[str]:
    """
    Ingests a document into ChromaDB.
    """
    try:
        ext = os.path.splitext(file_path)[1].lower()
        documents = []

        if ext == '.pdf':
            loader = PyPDFLoader(file_path)
            documents = loader.load()
        elif ext == '.docx':
            documents = load_docx(file_path)
        elif ext == '.txt':
            loader = TextLoader(file_path, encoding='utf-8')
            documents = loader.load()
        else:
            return []

        # OPTIMIZED LOGIC: 
        # 1. Use larger chunks (~2000 chars) for Nomic v1.5
        # 2. Add contextual header to every chunk
        
        # Prepend context to every page content BEFORE splitting
        # Use specific Nomic prefix if needed, but here we stick to the user's requested format.
        context_header = f"Document: {original_filename}\nSummary: {summary}\nContent: "
        
        for doc in documents:
            # We add headers to the content itself so the embedding captures it
            doc.page_content = f"{context_header}{doc.page_content}"
            doc.metadata['summary'] = summary
            doc.metadata['filename'] = original_filename

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2000, 
            chunk_overlap=400,
            separators=["\n\n", "Section", "ARTICLE", ".", " "] # Prioritize legal boundaries
        )
        chunks = text_splitter.split_documents(documents)
        
        for i, chunk in enumerate(chunks):
            chunk.metadata['filename'] = original_filename
            chunk.metadata['chunk_index'] = i
            chunk.metadata['source'] = original_filename

        if not chunks:
            return []

        # Use persistent DB
        # Distribute chunks across shards
        dbs = core.get_dbs()
        shard_chunks = [[] for _ in range(core.NUM_SHARDS)]
        
        for i, chunk in enumerate(chunks):
            shard_idx = i % core.NUM_SHARDS
            shard_chunks[shard_idx].append(chunk)
            
        all_ids = []
        all_ids = []
        BATCH_SIZE = 32

        for i, db_chunks in enumerate(shard_chunks):
            if db_chunks:
                total_chunks = len(db_chunks)
                print(f"Ingesting {total_chunks} chunks to shard {i}...")
                
                for j in range(0, total_chunks, BATCH_SIZE):
                    batch = db_chunks[j : j + BATCH_SIZE]
                    try:
                        ids = dbs[i].add_documents(batch)
                        all_ids.extend(ids)
                        print(f"  Processed batch {j//BATCH_SIZE + 1} ({len(batch)} chunks)")
                    except Exception as batch_error:
                        print(f"  Error processing batch {j//BATCH_SIZE + 1}: {batch_error}")
        
        print(f"Ingested total {len(all_ids)} chunks for {original_filename}")
        
        # Invalidate cache on new ingest
        core.clear_cache()
        
        return all_ids

    except Exception as e:
        print(f"Error during ingestion: {e}")
        return []

def generate_summary(file_path: str) -> str:
    """
    Generate summary using smart document processor with OCR fallback
    
    This now uses the document_processor module which implements:
    - Standard text extraction
    - OCR fallback for scanned documents
    - Qwen-based summary generation
    """
    try:
        # Use the smart document processor
        result = document_processor.process_document_complete(file_path, os.path.basename(file_path))
        
        if result.get("success"):
            return result.get("summary", "Summary generation completed.")
        else:
            error_msg = result.get("error", "Unknown error")
            print(f"Error in document processing: {error_msg}")
            return f"Could not generate summary: {error_msg}"
    
    except Exception as e:
        print(f"Error generating summary: {e}")
        import traceback
        traceback.print_exc()
        return "Summary generation failed."



def smart_ingest_document(
    file_path: str, 
    original_filename: str, 
    status_callback=None
) -> dict:
    """
    Smart ingestion using the Master Pipeline.
    
    This function:
    1. Auto-detects if the document is a Statute (Act) or Judgment (Case)
    2. Routes to the appropriate processor:
       - STATUTE → Section-based chunking with legal metadata
       - JUDGMENT → Zone-based chunking with ratio extraction
    3. Stores in dual-brain system (Vector DB + JSON Map)
    
    Returns:
        {
            "success": bool,
            "doc_type": "STATUTE" | "JUDGMENT",
            "chunks_created": int,
            "file_id": str,
            "message": str
        }
    """
    
    try:
        import master_ingest
        
        print(f"🚀 Starting Master Ingestion Pipeline for {original_filename}...")
        
        # Use the master pipeline
        result = master_ingest.smart_ingest(
            file_path=file_path,
            filename=original_filename,
            status_callback=status_callback
        )
        
        return result
    
    except Exception as e:
        print(f"❌ Master ingestion failed: {e}")
        import traceback
        traceback.print_exc()
        
        # Fallback to legacy processing
        print(f"⚠️  Falling back to legacy ingestion...")
        try:
            summary = generate_summary(file_path)
            chunk_ids = ingest_document(file_path, original_filename, summary)
            
            return {
                "success": len(chunk_ids) > 0,
                "doc_type": "UNKNOWN",
                "chunks_created": len(chunk_ids),
                "file_id": original_filename,
                "message": f"Processed {len(chunk_ids)} chunks using legacy mode (fallback)",
                "processing_method": "legacy_fallback"
            }
        except Exception as fallback_error:
            return {
                "success": False,
                "doc_type": "ERROR",
                "error": str(fallback_error),
                "message": "Both master and fallback processing failed"
            }


